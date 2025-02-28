from fastapi import FastAPI, Form, HTTPException
from fastapi.responses import FileResponse
from docx import Document
import json
import os
import subprocess
import requests

app = FastAPI()

# URL del archivo DOCX en S3
S3_DOCX_URL = "https://caribbeangoods-content-s3.s3.eu-west-2.amazonaws.com/CoffeeAgreement.docx"

# Rutas temporales
TEMP_DOCX = "CoffeeAgreement.docx"
OUTPUT_PDF = "CoffeeAgreement.pdf"

def download_file_from_s3(url, save_path):
    """ Descarga el archivo desde S3 y lo guarda localmente """
    try:
        response = requests.get(url, stream=True)
        response.raise_for_status()  # Lanza error si la respuesta no es 200
        with open(save_path, "wb") as file:
            for chunk in response.iter_content(chunk_size=8192):
                file.write(chunk)
        print(f"✅ Archivo descargado desde S3: {save_path}")
    except requests.exceptions.RequestException as e:
        print(f"❌ Error al descargar archivo desde S3: {e}")
        raise HTTPException(status_code=500, detail=f"Error al descargar archivo desde S3: {e}")

def convert_docx_to_pdf(docx_path, pdf_path):
    """ Convierte un archivo DOCX a PDF usando LibreOffice """
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path],
            check=True
        )
        print(f"✅ DOCX convertido a PDF correctamente: {pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"❌ Error al convertir DOCX a PDF: {e}")
        raise HTTPException(status_code=500, detail="Error al convertir DOCX a PDF")

@app.post("/edit-docx-to-pdf/")
async def edit_docx_to_pdf(replacements: str = Form(...)):
    """ Descarga el DOCX desde S3, reemplaza valores y lo convierte a PDF """
    
    try:
        replacements = json.loads(replacements)  # Convertir JSON a diccionario
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON format in replacements")

    # 📌 **Descargar el archivo desde S3**
    download_file_from_s3(S3_DOCX_URL, TEMP_DOCX)

    # 📌 **Abrir y modificar el DOCX**
    doc = Document(TEMP_DOCX)

    # 📌 **Aplicar reemplazos en párrafos**
    for para in doc.paragraphs:
        texto_completo = "".join(run.text for run in para.runs)
        for clave, valor in replacements.items():
            if clave in texto_completo:
                texto_completo = texto_completo.replace(clave, valor)
                para.clear()
                para.add_run(texto_completo)

    # 📌 **Aplicar reemplazos en tablas**
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    texto_completo = "".join(run.text for run in para.runs)
                    for clave, valor in replacements.items():
                        if clave in texto_completo:
                            texto_completo = texto_completo.replace(clave, valor)
                            para.clear()
                            para.add_run(texto_completo)

    # 📌 **Guardar el documento editado**
    doc.save(TEMP_DOCX)

    # 📌 **Convertir a PDF**
    convert_docx_to_pdf(TEMP_DOCX, OUTPUT_PDF)

    # 📌 **Devolver el archivo PDF generado**
    return FileResponse(OUTPUT_PDF, media_type="application/pdf", filename="CoffeeAgreement.pdf")
